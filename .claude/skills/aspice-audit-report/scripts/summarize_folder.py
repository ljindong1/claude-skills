#!/usr/bin/env python3
"""대상 산출물 폴더의 각 파일을 '내용 기반'으로 압축 요약(다이제스트)한다.

목적: AI가 이 다이제스트 하나를 읽고 파일↔체크리스트 산출물 매핑 테이블을
내용 근거로 생성하게 하기 위함(파일명 키워드 매칭의 한계 극복).
전체 문서를 읽지 않고 시트명/헤딩/첫 행/검출 ID/버전만 뽑아 토큰을 아낀다.

usage: python summarize_folder.py <폴더> [--out folder_digest.md]
"""
import os
import re
import sys
import glob
import argparse
import zipfile

ID = re.compile(r'\[([A-Za-z]{2,}[A-Za-z]?)-\d+\]')


def ids_in(texts):
    s = set()
    for t in texts:
        for m in ID.finditer(str(t)):
            s.add(m.group(1))
    return sorted(s)


def sum_xlsx(path):
    import openpyxl
    wb = openpyxl.load_workbook(path, read_only=True, data_only=True)
    sheets = wb.sheetnames
    first = wb[sheets[0]]
    rows = []
    for i, r in enumerate(first.iter_rows(min_row=1, max_row=8, values_only=True)):
        cs = [str(c)[:20] for c in r if c is not None][:6]
        if cs:
            rows.append(' | '.join(cs))
        if len(rows) >= 4:
            break
    return sheets, rows, ids_in(rows)


def sum_xls(path):
    import xlrd
    wb = xlrd.open_workbook(path)
    sheets = wb.sheet_names()
    sh = wb.sheet_by_index(0)
    rows = []
    for ri in range(min(sh.nrows, 8)):
        cs = [str(sh.cell_value(ri, ci))[:20] for ci in range(min(sh.ncols, 6))
              if str(sh.cell_value(ri, ci)).strip()]
        if cs:
            rows.append(' | '.join(cs))
        if len(rows) >= 4:
            break
    return sheets, rows, ids_in(rows)


def sum_docx(path):
    import docx
    d = docx.Document(path)
    heads = [p.text.strip() for p in d.paragraphs
             if p.style.name.startswith('Heading') and p.text.strip()][:12]
    return heads, len(d.tables)


def sum_dir(path):
    """하위 폴더를 한 산출물 단위로 요약 — 내부 파일 목록 + 대표 docx 헤딩."""
    inner = []
    for root, _, fs in os.walk(path):
        for fn in fs:
            inner.append(os.path.relpath(os.path.join(root, fn), path))
    kinds = {}
    for fn in inner:
        e = os.path.splitext(fn)[1].lower()
        kinds[e] = kinds.get(e, 0) + 1
    heads = []
    first_docx = next((os.path.join(path, x) for x in inner if x.lower().endswith('.docx')), None)
    if first_docx:
        try:
            import docx
            d = docx.Document(first_docx)
            heads = [pp.text.strip() for pp in d.paragraphs
                     if pp.style.name.startswith('Heading') and pp.text.strip()][:6]
        except Exception:
            pass
    return inner, kinds, heads


def extract_zips(folder, workdir):
    """폴더(+1단계 하위)의 zip을 workdir/<zip이름>/에 해제. 원본 폴더는 건드리지 않는다.

    중복 방지: 산출물 폴더에 zip과 같은 이름의 폴더가 이미 있으면(사용자가 직접 푼 경우)
    해제를 건너뛰고 그 폴더를 쓴다. 반환: [(zip경로, 해제폴더 또는 None(스킵), 사유)]
    """
    done = []
    cand = []
    for root in [folder] + [os.path.join(folder, d) for d in sorted(os.listdir(folder))
                            if os.path.isdir(os.path.join(folder, d)) and not d.startswith('_')]:
        for f in sorted(os.listdir(root)):
            if f.lower().endswith('.zip') and os.path.isfile(os.path.join(root, f)):
                cand.append(os.path.join(root, f))
    for zp in cand:
        stem = os.path.splitext(os.path.basename(zp))[0]
        sibling = os.path.join(os.path.dirname(zp), stem)
        if os.path.isdir(sibling):
            done.append((zp, None, f"동일 이름 폴더 존재({stem}/) — 이미 해제됨, 건너뜀"))
            continue
        dest = os.path.join(workdir, stem)
        try:
            os.makedirs(dest, exist_ok=True)
            with zipfile.ZipFile(zp) as z:
                z.extractall(dest)
            done.append((zp, dest, "해제 완료"))
        except Exception as e:
            done.append((zp, None, f"해제 실패: {str(e)[:40]} — 수동 해제 필요(암호/형식)"))
    return done


def main(folder, out, extract_dir=None):
    entries = sorted(os.listdir(folder))
    files = [f for f in entries if os.path.isfile(os.path.join(folder, f))]
    # `_`로 시작하는 폴더(작업 산출물: _오디트결과/_압축해제/_기준양식 등)는 점검대상에서 제외
    dirs = [f for f in entries if os.path.isdir(os.path.join(folder, f)) and not f.startswith('_')]
    extracted = extract_zips(folder, extract_dir) if extract_dir else []
    lines = [f"# 폴더 다이제스트 — {os.path.basename(folder)} (파일 {len(files)} + 폴더 {len(dirs)})",
             "", "각 항목의 내용 요약. AI가 이걸 읽고 체크리스트 산출물에 매핑한다.",
             "폴더 항목(📁)은 그 폴더 전체가 하나의 산출물(예: 모듈별로 나뉜 상세설계서)일 수 있다.", ""]
    for d in dirs:
        dp = os.path.join(folder, d)
        inner, kinds, heads = sum_dir(dp)
        lines.append(f"## 📁 {d}/  (하위 파일 {len(inner)}개)")
        lines.append(f"- 구성: {kinds}")
        lines.append(f"- 내부 파일(일부): {[os.path.basename(x) for x in inner[:10]]}")
        if heads:
            lines.append(f"- 대표 문서 헤딩: {heads}")
        lines.append("")
    for zp, dest, msg in extracted:
        lines.append(f"## 🗜 {os.path.basename(zp)} → {msg}")
        if dest:
            inner, kinds, heads = sum_dir(dest)
            lines.append(f"- 해제 위치: {dest}")
            lines.append(f"- 구성: {kinds}")
            lines.append(f"- 내부 파일(일부): {[os.path.basename(x) for x in inner[:10]]}")
            if heads:
                lines.append(f"- 대표 문서 헤딩: {heads}")
        lines.append("")
    for f in files:
        p = os.path.join(folder, f)
        ext = os.path.splitext(f)[1].lower()
        kb = os.path.getsize(p) // 1024
        lines.append(f"## {f}  ({kb}KB)")
        try:
            if ext == '.xlsx':
                sh, rows, ids = sum_xlsx(p)
                lines.append(f"- 시트: {sh[:6]}")
                if rows: lines.append(f"- 첫행: {rows[0][:80]}")
                if ids: lines.append(f"- 검출 ID: {ids}")
            elif ext == '.xls':
                sh, rows, ids = sum_xls(p)
                lines.append(f"- 시트: {sh[:6]}")
                if rows: lines.append(f"- 첫행: {rows[0][:80]}")
                if ids: lines.append(f"- 검출 ID: {ids}")
            elif ext == '.docx':
                heads, nt = sum_docx(p)
                lines.append(f"- 헤딩: {heads}")
                lines.append(f"- 표 {nt}개")
            elif ext == '.zip':
                with zipfile.ZipFile(p) as z:
                    inner = [n for n in z.namelist() if not n.endswith('/')][:8]
                lines.append(f"- 압축 내용: {inner}")
                if not extract_dir:
                    lines.append("- (내용 정독하려면 --extract-zips <작업폴더> 로 자동 해제)")
            elif ext == '.pdf':
                lines.append("- (PDF — 파일명 기준, 내용요약 생략)")
            else:
                lines.append(f"- ({ext} — 요약 미지원)")
        except Exception as e:
            lines.append(f"- 요약 실패: {str(e)[:50]}")
        lines.append("")
    text = '\n'.join(lines)
    if out:
        open(out, 'w', encoding='utf-8').write(text)
        print("저장:", out, f"({len(files)}개 파일)")
    else:
        print(text)


if __name__ == '__main__':
    import argparse
    ap = argparse.ArgumentParser()
    ap.add_argument('folder')
    ap.add_argument('--out', default='folder_digest.md')
    ap.add_argument('--extract-zips', dest='extract_dir', default=None,
                    help='zip을 이 작업 폴더에 자동 해제해 다이제스트에 포함(원본 폴더 불변)')
    a = ap.parse_args()
    main(a.folder, a.out, a.extract_dir)
