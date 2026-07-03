#!/usr/bin/env python3
"""산출물(docx/xlsx)의 전체 본문을 '절 구조'까지 보존해 추출한다.

목적: AI가 사람처럼 코멘트를 쓰려면 헤딩 요약이 아니라 본문 전체(절 번호·필드·표 내용)를
읽어야 한다. 이 추출 결과를 AI가 읽고 각 체크리스트 항목의 근거 위치(절/필드)를 인용한다.

usage: python extract_deliverable_text.py <파일> [--out text.md] [--max-tables N]
"""
import sys
import os
import argparse


def from_docx(path, max_tables):
    import docx
    d = docx.Document(path)
    out = []
    for p in d.paragraphs:
        t = p.text.strip()
        if not t:
            continue
        style = p.style.name
        if style.startswith("Heading"):
            lvl = style.replace("Heading", "").strip()
            lvl = int(lvl) if lvl.isdigit() else 1
            out.append("\n" + "#" * min(lvl + 1, 6) + " " + t)
        else:
            out.append(t)
    out.append("\n--- 표(Tables) ---")
    for i, tb in enumerate(d.tables[:max_tables], 1):
        out.append(f"\n[표 {i}]")
        for r in tb.rows[:30]:
            cells = [c.text.strip().replace("\n", " ") for c in r.cells]
            cells = [c for c in cells if c]
            if cells:
                out.append(" | ".join(cells)[:300])
    if len(d.tables) > max_tables:
        out.append(f"\n... (표 {len(d.tables)}개 중 {max_tables}개만)")
    return "\n".join(out)


def from_xlsx(path):
    import openpyxl
    wb = openpyxl.load_workbook(path, read_only=True, data_only=True)
    out = []
    for sh in wb.sheetnames:
        ws = wb[sh]
        out.append(f"\n## 시트: {sh}")
        for r in ws.iter_rows(min_row=1, max_row=60, values_only=True):
            cells = [str(c).strip().replace("\n", " ") for c in r if c is not None]
            if cells:
                out.append(" | ".join(cells)[:300])
    return "\n".join(out)


def main(path, out, max_tables):
    ext = os.path.splitext(path)[1].lower()
    if ext == ".docx":
        text = from_docx(path, max_tables)
    elif ext in (".xlsx", ".xlsm"):
        text = from_xlsx(path)
    else:
        text = f"(지원하지 않는 형식: {ext} — .xls는 별도 변환 필요)"
    head = f"# 본문 추출 — {os.path.basename(path)}\n"
    if out:
        open(out, "w", encoding="utf-8").write(head + text)
        print("저장:", out, f"({len(text)}자)")
    else:
        print(head + text)


if __name__ == "__main__":
    ap = argparse.ArgumentParser()
    ap.add_argument("path")
    ap.add_argument("--out", default=None)
    ap.add_argument("--max-tables", type=int, default=40)
    a = ap.parse_args()
    main(a.path, a.out, a.max_tables)
